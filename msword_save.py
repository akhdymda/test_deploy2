from docx import Document # ワードファイルを生成するための機能をインポート
import io

# 要約をワードに保存するための関数の設定
def save_summary_to_word(summary, file_name='summary.docx'):
    document = Document() # 空のドキュメントファイルを作成
    document.add_heading('生成された要約', level=1) # 表題を追加
    document.add_paragraph(summary) # テキストを追加
    # メモリ上にWordファイルを保存
    byte_io = io.BytesIO()
    document.save(byte_io)
    byte_io.seek(0)  # ファイルポインタを先頭に戻す
    return byte_io